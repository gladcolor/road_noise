from PIL import Image
import os
import glob
import collections

import matplotlib.pyplot as plt

import torch
import torchvision
import torch.nn as nn
import numpy as np
from torchvision import datasets, models, transforms
import shutil
import pandas as pd

import docx
import glob
import os
from copy import deepcopy
import sys
import logging
import yaml
import logging.config



def setup_logging(default_path='log_config.yaml', logName='', default_level=logging.DEBUG):
    path = default_path
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
            if logName !='':
                config["handlers"]["file"]['filename'] = logName
            logging.config.dictConfig(config)
    else:
        logging.basicConfig(level=default_level)

log_txt_name = r'K:\Research\Tagme\tagme\src\assets\log.log'
yaml_path = 'log_config.yaml'
setup_logging(yaml_path)
        # logger.info(os.path.basename(file))

logger = logging.getLogger('main.core')

class ImageClassifier():

    def __init__(self, model_path, input_size=299):

        """ Inception v3
        Be careful, expects (299,299) sized images and has auxiliary output
        """
        self.device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")
        self.device = 'cpu'

        # model_ft = models.inception_v3(pretrained=None)
        # num_classes = 2
        #
        # num_ftrs = model_ft.AuxLogits.fc.in_features
        # model_ft.AuxLogits.fc = nn.Linear(num_ftrs, num_classes)
        # # Handle the primary net
        # num_ftrs = model_ft.fc.in_features
        # model_ft.fc = nn.Linear(num_ftrs, num_classes)
        #
        #
        # model = model_ft

        # self.model = model
        self.model = torch.load(model_path)
        self.model.eval()

        #
        # if model_path is not None:
        #     self.model_path = model_path
        #     # self.model = torch.load(model_path)
        #     # self.model.load_state_dict(torch.load(self.model_path))
        #     self.model = torch.load(self.model_path)
        #     self.model.eval()

        self.model = self.model.to(self.device)
        self.input_size = input_size
    #
    # def getModel(self, model_path):
    #     model = torch.load(model_path)
    #     return model

    def image_inference(self, image_path):
        img = Image.open(image_path)
        img = transforms.Resize(self.input_size)(img)
        img = transforms.ToTensor()(img)
        img = torch.unsqueeze(img, 0)
        img = img.to(self.device)
        result = self.model(img)
        logger.info(os.path.basename(image_path))
        logger.info(result.cpu().detach().numpy())
        result = nn.Softmax()(result)

        return result.cpu().detach().numpy()

def inference():
    model_path = r'inception.pth'
    classifier = ImageClassifier(model_path=model_path)
    # test_image_path = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2\hlvA4Deozc_zh2J-gDUVUg_-70.975509_42.372362_0_82.00.jpg'
    # images = glob.glob(r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2\*.jpg')
    df = pd.read_csv(r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\val.csv')
    images = []
    for idx, row in df.iterrows():
        images.append(row['image'])

    copy_to = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\label1'
    print(f'Images counts: {len(images)}')
    for jpg in images:
        result = classifier.image_inference(jpg)
        basename = os.path.basename(jpg)
        label = np.argmax(result)
        # print(f"{basename} | label: {label}")
        if label == 1:
            new_name = os.path.join(copy_to, basename)
            shutil.copyfile(jpg, new_name)
            print(f"{basename}")


# https://www.zhihu.com/question/68384370 通过pytorch的hook机制简单实现了一下，只输出conv层的特征图。

def viz(module, input):
    x = input[0][0]
    #最多显示4张图
    min_num = np.minimum(4, x.size()[0])
    for i in range(min_num):
        plt.subplot(1, 4, i+1)
        plt.imshow(x[i])
    plt.show()

def viz2(module, input):
    print("type of input:", type(input))
    print("len of input:", len(input))
    print("Hook input:", input)

def getInceptionFeature():
    model_path = r'inception_all.pth'

    model = torch.load(model_path)
    model.eval()

    for name, m in model.named_modules():
        if isinstance(m, torch.nn.Linear):
            m.register_forward_pre_hook(viz2)


    print(model)

    features_model = torch.nn.Sequential(collections.OrderedDict(list(model.named_children())[:-1]))
    print(features_model)
    features_model.eval()

    # test_image_path = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2\hlvA4Deozc_zh2J-gDUVUg_-70.975509_42.372362_0_82.00.jpg'
    # images = glob.glob(r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2\*.jpg')
    df = pd.read_csv(r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\val.csv')
    img_dir = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2'
    images = []
    device = 'cpu'
    input_size = 299
    for idx, row in df.iterrows():
        images.append(os.path.join(img_dir, row['image']))

    copy_to = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\label1'
    print(f'Images counts: {len(images)}')
    for jpg in images:
        basename = os.path.basename(jpg)

        img = Image.open(jpg)
        img = transforms.Resize(input_size)(img)
        img = transforms.ToTensor()(img)
        img = torch.unsqueeze(img, 0)
        img = img.to(device)

        with torch.no_grad():
            result = model(img)
        # result = features_model(img)
        # result = model(img)

        logger.info(os.path.basename(jpg))
        logger.info(result.cpu().detach().numpy())



def inference_to_docx():
    img_dir = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\random100'
    model_path = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\project_code\inception_all.pth'
    classifier = ImageClassifier(model_path=model_path)

    # docx operation
    template = docx.Document(r'K:\Research\Resilience\template.docx')
    t_table = template.tables[0]._tbl
    document = docx.Document()

    csv_file = r'K:\Research\Tagme\tagme\src\assets\val.csv'
    predict_file = csv_file.replace(".csv", "_inferenced.csv")
    log_txt_name = csv_file.replace('.csv', '.log')
    # logger = logging.getLogger(log_txt_name)
    # test_image_path = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2\hlvA4Deozc_zh2J-gDUVUg_-70.975509_42.372362_0_82.00.jpg'
    # images = glob.glob(r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\building_images2\*.jpg')
    df = pd.read_csv(csv_file)
    images = []
    truths = []
    for idx, row in df[:10].iterrows():
        images.append(row['image'])
        truths.append(row['class_id'])

    yaml_path = 'log_config.yaml'
    setup_logging(yaml_path, logName=log_txt_name)
    # logger.info(os.path.basename(file))

    # copy_to = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\label1'
    print(f'Images counts: {len(images)}')
    f = open(predict_file, 'w')
    f.writelines("image,predict,label\n")
    for idx, jpg in enumerate(images[:]):
        jpg = os.path.join(img_dir, jpg)

        result = classifier.image_inference(jpg)
        basename = os.path.basename(jpg)
        predict = np.argmax(result)
        info = f"{basename} | predict: {predict}"
        logger.info(info)
        # print()

        # docx operation
        new_tbl = deepcopy(t_table)
        paragraph = document.add_paragraph()
        paragraph._p.addnext(new_tbl)

        new_tbl = document.tables[idx]

        new_tbl.rows[0].cells[0].text = basename

        #     row_cells = tbl.add_row().cells
        paragraph = new_tbl.rows[1].cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(jpg, width=2500000, height=3300000)

        # fill prediction:
        if predict == 1:
            new_tbl.rows[1].cells[2].text = "resilience"
        else:
            new_tbl.rows[1].cells[2].text = "non_resilience"

        # fill the truth:
        truth = truths[idx]
        if truth == 1:
            new_tbl.rows[2].cells[2].text = "resilience"
        else:
            new_tbl.rows[2].cells[2].text = "non_resilience"

        # two tabels per page
        table_cnt = idx + 1
        if (table_cnt % 2 == 0) and (table_cnt > 0):
            document.add_page_break()

        f.writelines(f'{jpg},{predict},{truth}\n')

    document.save(r"K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\random100\reference_results.docx")

def inference_folder_to_docx():
        img_dir = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\random100'
        model_path = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\project_code\inception_all.pth'
        classifier = ImageClassifier(model_path=model_path)

        # docx operation
        template = docx.Document(r'K:\Research\Resilience\template.docx')
        t_table = template.tables[0]._tbl
        document = docx.Document()


        predict_file = os.path.join(img_dir, "inferenced.csv")
        log_txt_name = os.path.join(img_dir,  'info.log')


        images = glob.glob(os.path.join(img_dir, "*.jpg"))

        yaml_path = 'log_config.yaml'
        setup_logging(yaml_path, logName=log_txt_name)
        # logger.info(os.path.basename(file))

        # copy_to = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\label1'
        print(f'Images counts: {len(images)}')
        f = open(predict_file, 'w')
        f.writelines("image,predict,label\n")
        for idx, jpg in enumerate(images[:]):
            jpg = os.path.join(img_dir, jpg)

            result = classifier.image_inference(jpg)
            basename = os.path.basename(jpg)
            predict = np.argmax(result)
            info = f"{basename} | predict: {predict}"
            logger.info(info)
            # print()

            # docx operation
            new_tbl = deepcopy(t_table)
            paragraph = document.add_paragraph()
            paragraph._p.addnext(new_tbl)

            new_tbl = document.tables[idx]

            new_tbl.rows[0].cells[0].text = basename

            #     row_cells = tbl.add_row().cells
            paragraph = new_tbl.rows[1].cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(jpg, width=2500000, height=3300000)

            # fill prediction:
            if predict == 1:
                new_tbl.rows[1].cells[2].text = "resilience"
            else:
                new_tbl.rows[1].cells[2].text = "non_resilience"

            # two tabels per page
            table_cnt = idx + 1
            if (table_cnt % 2 == 0) and (table_cnt > 0):
                document.add_page_break()

            f.writelines(f'{jpg},{predict}\n')
        saved_docx = os.path.join(img_dir, "inferenced.docx")
        document.save(saved_docx)


def results_to_docx():
    img_dir = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\random100'
    result_file = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\random100\sample ys.csv'

    df = pd.read_csv(result_file)

    # docx operation
    template = docx.Document(r'K:\Research\Resilience\template.docx')
    t_table = template.tables[0]._tbl
    document = docx.Document()


    log_txt_name = os.path.join(img_dir, 'info.log')


    yaml_path = 'log_config.yaml'
    setup_logging(yaml_path, logName=log_txt_name)
    # logger.info(os.path.basename(file))

    # copy_to = r'K:\OneDrive_NJIT\OneDrive - NJIT\Research\Resilience\data\boston\label1'
    print(f'CSV row: {len(df)}')

    for idx, row in df.iterrows():
        try:
            jpg = row['imageUrl']
            basename = os.path.basename(jpg)
            jpg = os.path.join(img_dir, basename)


            info = f"{basename} "
            logger.info(info)
            # print()

            # docx operation
            new_tbl = deepcopy(t_table)
            paragraph = document.add_paragraph()
            paragraph._p.addnext(new_tbl)

            new_tbl = document.tables[idx]

            new_tbl.rows[0].cells[0].text = basename

            #     row_cells = tbl.add_row().cells
            paragraph = new_tbl.rows[1].cells[0].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(jpg, width=2500000, height=3300000)

            # fill the label:
            keyTag = row['keyTag']
            notes = row['notes']
            new_tbl.rows[2].cells[2].text = keyTag
            new_tbl.rows[4].cells[2].text = notes

            # two tabels per page
            table_cnt = idx + 1
            if (table_cnt % 2 == 0) and (table_cnt > 0):
                document.add_page_break()

        except Exception as e:
            logger.error(e)
            continue

    saved_docx = os.path.join(img_dir, "results.docx")
    document.save(saved_docx)

if __name__ == "__main__":
    # inference_to_docx();
    # inference_folder_to_docx()

    # results_to_docx()

    getInceptionFeature()